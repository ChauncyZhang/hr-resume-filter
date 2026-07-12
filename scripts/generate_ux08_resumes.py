from __future__ import annotations

import argparse
import csv
import json
import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "design" / "test-data" / "ux-08-resumes"
FILES_DIR = OUTPUT_DIR / "files"


def fixture(
    fixture_id: str,
    name: str,
    position: str,
    file_format: str,
    score: str,
    parse_status: str = "success",
    tags: tuple[str, ...] = (),
    skills: tuple[str, ...] = (),
    missing: str = "无明显缺失",
    email_suffix: str | None = None,
) -> dict:
    extension = "pdf" if file_format in {"pdf", "scan_pdf", "corrupt_pdf"} else file_format
    return {
        "id": fixture_id,
        "name": name,
        "filename": f"{position}_{name}_{fixture_id}.{extension}",
        "format": file_format,
        "targetPosition": position,
        "expectedParseStatus": parse_status,
        "expectedScoreRange": score,
        "scenarioTags": list(tags),
        "skills": list(skills),
        "missing": missing,
        "synthetic": True,
        "email": f"{email_suffix or fixture_id.lower()}@example.com",
        "phone": f"138****{fixture_id[-4:]}",
    }


FIXTURES = [
    fixture("SYN-0001", "林启舟", "AI 工程师", "pdf", "85-92", tags=("高匹配",), skills=("Python", "RAG", "Agent", "PyTorch")),
    fixture("SYN-0002", "苏若岚", "AI 工程师", "docx", "72-80", tags=("中等匹配",), skills=("Python", "机器学习", "TensorFlow"), missing="Agent 生产经验"),
    fixture("SYN-0003", "周宁", "AI 工程师", "txt", "80-88", tags=("同名不同人",), skills=("NLP", "LLM", "Prompt Engineering"), email_suffix="zhou.ning.ai"),
    fixture("SYN-0004", "顾言川", "AI 工程师", "scan_pdf", "68-76", tags=("扫描版", "缺少教育信息"), skills=("Python", "OCR", "OpenCV"), missing="教育经历"),
    fixture("SYN-0005", "许知远", "AI 工程师", "pdf", "48-58", tags=("低匹配",), skills=("数据分析", "SQL"), missing="大模型项目经验"),
    fixture("SYN-0006", "林启舟", "AI 工程师", "docx", "85-92", tags=("重复候选人", "新版简历"), skills=("Python", "RAG", "Agent", "PyTorch"), email_suffix="syn-0001"),
    fixture("SYN-0007", "程星野", "AI 工程师", "pdf", "70-78", tags=("缺少量化结果",), skills=("Python", "LangChain", "向量数据库"), missing="项目量化结果"),
    fixture("SYN-0008", "唐予安", "AI 工程师", "corrupt_pdf", "不可评分", parse_status="failed", tags=("损坏文件",), skills=()),
    fixture("SYN-0009", "陆谨言", "Java 后端工程师", "pdf", "84-90", tags=("高匹配",), skills=("Java", "Spring Boot", "MySQL", "Redis")),
    fixture("SYN-0010", "沈嘉禾", "Java 后端工程师", "docx", "70-78", tags=("中等匹配",), skills=("Java", "Spring Cloud", "MySQL"), missing="高并发量化结果"),
    fixture("SYN-0011", "方景行", "Java 后端工程师", "txt", "50-60", tags=("低匹配",), skills=("Java", "Android"), missing="服务端项目经验"),
    fixture("SYN-0012", "贺清越", "Java 后端工程师", "pdf", "76-84", tags=("缺少教育信息",), skills=("Java", "Kafka", "Kubernetes"), missing="教育经历"),
    fixture("SYN-0013", "周宁", "产品经理", "pdf", "78-86", tags=("同名不同人",), skills=("B 端产品", "需求分析", "项目管理"), email_suffix="zhou.ning.pm"),
    fixture("SYN-0014", "叶舒然", "产品经理", "docx", "68-76", tags=("中等匹配",), skills=("用户研究", "产品规划"), missing="AI 产品经验"),
    fixture("SYN-0015", "孟书瑶", "产品经理", "scan_pdf", "45-55", tags=("扫描版", "低匹配"), skills=("活动运营",), missing="产品完整交付经验"),
    fixture("SYN-0016", "江砚秋", "前端工程师", "pdf", "82-89", tags=("高匹配",), skills=("React", "TypeScript", "CSS", "可访问性")),
    fixture("SYN-0017", "白屿川", "前端工程师", "docx", "74-82", tags=("超长文件",), skills=("Vue", "React", "工程化"), missing="数据可视化"),
    fixture("SYN-0018", "季南乔", "前端工程师", "txt", "52-62", tags=("低匹配",), skills=("HTML", "CSS"), missing="现代前端工程经验"),
]


POSITION_SUMMARIES = {
    "AI 工程师": "负责虚构企业知识检索、模型应用和智能工作流研发。",
    "Java 后端工程师": "负责虚构业务系统服务端架构、稳定性与性能优化。",
    "产品经理": "负责虚构企业服务产品的需求分析、规划和跨团队交付。",
    "前端工程师": "负责虚构管理平台的前端架构、体验和工程质量。",
}


def resume_sections(item: dict) -> list[tuple[str, list[str]]]:
    years = 3 + int(item["id"][-2:]) % 6
    skills = "、".join(item["skills"]) or "待解析"
    quantified = "将核心流程耗时降低 31%，并建立可观测指标。" if "缺少量化结果" not in item["scenarioTags"] else "参与核心流程优化与交付。"
    education = [] if "缺少教育信息" in item["scenarioTags"] else ["云岭大学（虚构） · 计算机相关专业 · 本科"]
    return [
        ("个人概况", [f"{years} 年相关经验。{POSITION_SUMMARIES[item['targetPosition']]}", "本简历为 UX-08 合成测试数据，不对应任何真实个人。"]),
        ("核心技能", [skills]),
        ("工作经历", [f"星河科技（虚构） · {item['targetPosition']} · 2021 至今", quantified, "与产品、测试和业务团队协作完成需求拆解、交付和复盘。"]),
        ("项目经历", ["企业协同平台（虚构项目）", "负责方案设计、关键模块实现、测试验证和上线跟踪。", f"待确认项：{item['missing']}。"]),
        ("教育经历", education or ["未提供"]),
    ]


def set_run_font(run, name: str = "Microsoft YaHei", size: float = 11, bold: bool = False, color: str = "222222") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color in (("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def write_docx(item: dict, path: Path) -> None:
    document = Document()
    configure_docx(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run(item["name"]), size=22, bold=True, color="173D57")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run(f"{item['targetPosition']} | {item['phone']} | {item['email']}"), size=10.5, color="5D7180")
    for heading, lines in resume_sections(item):
        document.add_heading(heading, level=1)
        for line in lines:
            paragraph = document.add_paragraph(style="Normal")
            set_run_font(paragraph.add_run(line))
    if "超长文件" in item["scenarioTags"]:
        document.add_page_break()
        for index in range(1, 15):
            document.add_heading(f"补充项目 {index}", level=2)
            paragraph = document.add_paragraph()
            set_run_font(paragraph.add_run(f"虚构项目 {index}：负责需求拆解、组件设计、自动化测试和发布复盘。所有内容均为合成数据。"))
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run(f"UX-08 SYNTHETIC RESUME | {item['id']}"), size=9, color="777777")
    document.save(path)


def register_pdf_font() -> str:
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        return "Helvetica"


def write_text_pdf(item: dict, path: Path) -> None:
    font_name = register_pdf_font()
    pdf = canvas.Canvas(str(path), pagesize=LETTER)
    width, height = LETTER
    y = height - 64
    pdf.setFont(font_name, 20)
    pdf.drawString(64, y, item["name"])
    y -= 24
    pdf.setFont(font_name, 10)
    pdf.drawString(64, y, f"{item['targetPosition']} | {item['phone']} | {item['email']}")
    y -= 30
    for heading, lines in resume_sections(item):
        pdf.setFont(font_name, 13)
        pdf.drawString(64, y, heading)
        y -= 20
        pdf.setFont(font_name, 10.5)
        for line in lines:
            for segment in [line[index:index + 42] for index in range(0, len(line), 42)]:
                pdf.drawString(74, y, segment)
                y -= 16
        y -= 8
    pdf.setFont(font_name, 8.5)
    pdf.drawCentredString(width / 2, 34, f"UX-08 SYNTHETIC RESUME | {item['id']}")
    pdf.save()


def chinese_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [Path("C:/Windows/Fonts/msyh.ttc"), Path("C:/Windows/Fonts/simhei.ttf")]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def write_scan_pdf(item: dict, path: Path) -> None:
    image = Image.new("RGB", (1275, 1650), "white")
    draw = ImageDraw.Draw(image)
    title_font = chinese_font(38)
    heading_font = chinese_font(25)
    body_font = chinese_font(20)
    y = 90
    draw.text((90, y), item["name"], font=title_font, fill="#173D57")
    y += 62
    draw.text((90, y), f"{item['targetPosition']} | {item['phone']} | {item['email']}", font=body_font, fill="#526876")
    y += 70
    for heading, lines in resume_sections(item):
        draw.text((90, y), heading, font=heading_font, fill="#2E74B5")
        y += 42
        for line in lines:
            for segment in [line[index:index + 38] for index in range(0, len(line), 38)]:
                draw.text((105, y), segment, font=body_font, fill="#222222")
                y += 34
        y += 22
    draw.text((90, 1570), f"UX-08 SYNTHETIC RESUME | {item['id']}", font=body_font, fill="#777777")
    image.save(path, "PDF", resolution=144.0)


def write_txt(item: dict, path: Path) -> None:
    lines = [item["name"], f"{item['targetPosition']} | {item['phone']} | {item['email']}", "", "[UX-08 合成测试简历]"]
    for heading, values in resume_sections(item):
        lines.extend(["", heading, *[f"- {value}" for value in values]])
    lines.append(f"\nFixture ID: {item['id']} | synthetic=true")
    path.write_text("\n".join(lines), encoding="utf-8")


def write_fixture(item: dict, path: Path) -> None:
    if item["format"] == "docx":
        write_docx(item, path)
    elif item["format"] == "pdf":
        write_text_pdf(item, path)
    elif item["format"] == "scan_pdf":
        write_scan_pdf(item, path)
    elif item["format"] == "txt":
        write_txt(item, path)
    elif item["format"] == "corrupt_pdf":
        path.write_bytes(b"%PDF-UX08-CORRUPT\nThis fixture is intentionally invalid.\n")
    else:
        raise ValueError(f"Unsupported fixture format: {item['format']}")


def write_metadata() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    FILES_DIR.mkdir(parents=True, exist_ok=True)
    (OUTPUT_DIR / "manifest.json").write_text(json.dumps(FIXTURES, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    with (OUTPUT_DIR / "expected-results.csv").open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle)
        writer.writerow(["编号", "文件名", "目标岗位", "预期解析状态", "预期匹配分区间", "场景标签", "合成数据"])
        for item in FIXTURES:
            writer.writerow([item["id"], item["filename"], item["targetPosition"], item["expectedParseStatus"], item["expectedScoreRange"], "、".join(item["scenarioTags"]), "是"])
    (OUTPUT_DIR / "README.md").write_text(
        "# UX-08 合成简历测试包\n\n"
        "本目录包含 18 份完全虚构的招聘流程测试简历。所有姓名、单位、学校、项目、邮箱和电话均为合成数据。\n\n"
        "- 仅用于 UX-08 原型、解析和可用性测试。\n"
        "- 不得替换为生产候选人简历。\n"
        "- 真实用户测试时只能使用本目录文件。\n"
        "- `manifest.json` 是程序清单，`expected-results.csv` 是中文人工对照表。\n"
        "- 损坏 PDF 为故意构造的错误样本，不应被正常解析。\n",
        encoding="utf-8",
    )


def generate_output() -> None:
    write_metadata()
    expected_names = {item["filename"] for item in FIXTURES}
    for existing in FILES_DIR.iterdir():
        if existing.name not in expected_names:
            existing.unlink()
    for item in FIXTURES:
        write_fixture(item, FILES_DIR / item["filename"])


def verify_output(output_dir: Path = OUTPUT_DIR) -> None:
    manifest_path = output_dir / "manifest.json"
    if not manifest_path.exists():
        raise FileNotFoundError(f"Missing manifest: {manifest_path}")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert len(manifest) == 18, "Expected exactly 18 fixture records"
    assert len({item["id"] for item in manifest}) == 18, "Fixture IDs must be unique"
    assert len({item["filename"] for item in manifest}) == 18, "Fixture filenames must be unique"
    assert all(item["synthetic"] is True for item in manifest), "Every fixture must be synthetic"
    assert Counter(item["targetPosition"] for item in manifest) == {
        "AI 工程师": 8,
        "Java 后端工程师": 4,
        "产品经理": 3,
        "前端工程师": 3,
    }
    tags = {tag for item in manifest for tag in item["scenarioTags"]}
    assert {"重复候选人", "同名不同人", "扫描版", "损坏文件", "超长文件"}.issubset(tags)
    for item in manifest:
        assert item["email"].endswith("@example.com")
        assert re.fullmatch(r"138\*{4}\d{4}", item["phone"])
        file_path = output_dir / "files" / item["filename"]
        assert file_path.exists() and file_path.stat().st_size > 20, f"Missing fixture file: {file_path}"
    print("18 synthetic resumes verified")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--verify-only", action="store_true")
    args = parser.parse_args()
    if not args.verify_only:
        generate_output()
    verify_output()


if __name__ == "__main__":
    main()
