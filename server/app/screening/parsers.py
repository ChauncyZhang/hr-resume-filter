import io
import math
import re
import zipfile
from collections import Counter
from dataclasses import dataclass
from typing import BinaryIO

from server.app.queue.service import normalize_safe_code
from server.app.resume_text import is_obfuscation_marker, normalize_resume_line, sanitize_resume_text
from server.app.screening.structured_pdf import StructuredPdfError, extract_structured_pdf

@dataclass(frozen=True)
class ParserLimits:
    max_source_bytes: int = 10 * 1024 * 1024; max_text_chars: int = 500_000; pdf_max_pages: int = 100
    docx_max_entries: int = 1000; docx_max_uncompressed_bytes: int = 50 * 1024 * 1024; docx_max_compression_ratio: int = 100

@dataclass(frozen=True)
class ParsedDocument:
    text: str; parser_version: str; quality: str

class ParserError(Exception):
    def __init__(self, safe_code: str) -> None: self.safe_code = normalize_safe_code(safe_code); super().__init__(self.safe_code)

_MIMES = {".pdf": "application/pdf", ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".txt": "text/plain"}

def parse_document(stream: BinaryIO, *, extension: str, mime_type: str, limits: ParserLimits = ParserLimits()) -> ParsedDocument:
    extension = extension.lower()
    if extension not in _MIMES: raise ParserError("file_type_not_allowed")
    if mime_type != _MIMES[extension]: raise ParserError("file_type_mismatch")
    stream.seek(0); data = stream.read(limits.max_source_bytes + 1)
    if len(data) > limits.max_source_bytes: raise ParserError("file_too_large")
    if extension == ".pdf":
        if not data.startswith(b"%PDF-"): raise ParserError("file_magic_mismatch")
        return _pdf(data, limits)
    if extension == ".docx":
        if not data.startswith(b"PK\x03\x04"): raise ParserError("file_magic_mismatch")
        return _docx(data, limits)
    if data.startswith((b"%PDF-", b"PK\x03\x04")): raise ParserError("file_magic_mismatch")
    return _txt(data, limits)

def validate_upload_preflight(stream: BinaryIO, *, extension: str, mime_type: str, limits: ParserLimits = ParserLimits()) -> str:
    extension=extension.lower()
    if extension not in _MIMES: raise ParserError("file_type_not_allowed")
    if mime_type!=_MIMES[extension]: raise ParserError("file_type_mismatch")
    stream.seek(0); data=stream.read(limits.max_source_bytes+1); stream.seek(0)
    if not data: raise ParserError("empty_file")
    if len(data)>limits.max_source_bytes: raise ParserError("file_too_large")
    if extension==".pdf" and not data.startswith(b"%PDF-"): raise ParserError("file_magic_mismatch")
    if extension==".docx":
        if not data.startswith(b"PK\x03\x04"): raise ParserError("file_magic_mismatch")
        _docx_preflight(data,limits)
    if extension==".txt":
        if data.startswith((b"%PDF-",b"PK\x03\x04")): raise ParserError("file_magic_mismatch")
        if b"\x00" in data: raise ParserError("binary_text_rejected")
    return extension[1:]

def _docx_preflight(data,limits):
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos=archive.infolist()
            if len(infos)>limits.docx_max_entries: raise ParserError("docx_entry_limit")
            total=0
            for info in infos:
                parts=info.filename.replace("\\","/").split("/")
                if ".." in parts or info.filename.startswith(("/","\\")): raise ParserError("docx_path_traversal")
                if info.filename.lower().endswith(("vbaproject.bin",".docm",".dotm")): raise ParserError("docx_macro_rejected")
                total+=info.file_size
                if total>limits.docx_max_uncompressed_bytes: raise ParserError("docx_size_limit")
                if info.file_size and info.file_size/max(1,info.compress_size)>limits.docx_max_compression_ratio: raise ParserError("docx_compression_ratio")
    except ParserError: raise
    except zipfile.BadZipFile: raise ParserError("docx_malformed") from None

def _bounded(text: str, limits: ParserLimits) -> str:
    if len(text) > limits.max_text_chars: raise ParserError("text_limit_exceeded")
    return text


def _pdf_page_text(page) -> str:
    fragments: list[dict[str, float | int | str]] = []

    def visitor(text, _cm, tm, _font, _size) -> None:
        value = normalize_resume_line(text)
        if value:
            fragments.append({"text": value, "x": float(tm[4]), "y": float(tm[5]), "order": len(fragments)})

    fallback = page.extract_text(visitor_text=visitor) or ""
    plain_text = sanitize_resume_text(fallback)
    page_box = getattr(page, "mediabox", None)
    page_width = float(getattr(page_box, "width", 600) or 600)
    page_height = float(getattr(page_box, "height", 800) or 800)
    outside = sum(
        float(fragment["x"]) < -12
        or float(fragment["x"]) > page_width * 1.08
        or float(fragment["y"]) < -12
        or float(fragment["y"]) > page_height * 1.08
        for fragment in fragments
    )
    origin_dates = sum(
        abs(float(fragment["x"])) <= 2
        and abs(float(fragment["y"])) <= 2
        and re.search(r"(?:19|20)\d{2}\s*年?\s*\d{0,2}", str(fragment["text"])) is not None
        for fragment in fragments
    )
    if fragments and (outside >= max(2, math.ceil(len(fragments) * 0.12)) or origin_dates >= 3):
        return plain_text

    marker_counts = Counter(str(fragment["text"]) for fragment in fragments if is_obfuscation_marker(str(fragment["text"])))
    repeated_markers = {value for value, count in marker_counts.items() if count >= 2}
    fragments = [fragment for fragment in fragments if str(fragment["text"]) not in repeated_markers]
    if not fragments:
        return sanitize_resume_text(fallback)

    rows: list[dict[str, object]] = []
    for fragment in sorted(fragments, key=lambda item: (float(item["y"]), int(item["order"]))):
        if rows and abs(float(rows[-1]["y"]) - float(fragment["y"])) <= 2:
            rows[-1]["parts"].append(fragment)
        else:
            rows.append({"y": fragment["y"], "parts": [fragment]})

    ordered_rows: list[dict[str, float | str]] = []
    for row in rows:
        parts = sorted(row["parts"], key=lambda item: (float(item["x"]), int(item["order"])))
        ordered_rows.append({
            "x": min(float(item["x"]) for item in parts),
            "y": float(row["y"]),
            "text": normalize_resume_line("".join(str(item["text"]) for item in parts)),
        })

    starts = sorted({round(float(row["x"]), 1) for row in ordered_rows if float(row["x"]) >= 0})
    split: float | None = None
    widest_gap = 0.0
    minimum_gap = page_width * 0.18
    for left, right in zip(starts, starts[1:]):
        gap = right - left
        left_count = sum(float(row["x"]) <= left for row in ordered_rows)
        right_count = sum(float(row["x"]) >= right for row in ordered_rows)
        if gap > minimum_gap and left_count >= 3 and right_count >= 3 and gap > widest_gap:
            split = (left + right) / 2
            widest_gap = gap

    ordered_rows.sort(key=lambda row: (
        0 if split is not None and float(row["x"]) < split else 1 if split is not None else 0,
        float(row["y"]),
    ))
    return sanitize_resume_text("\n".join(str(row["text"]) for row in ordered_rows if row["text"]))

def _pdf_fallback(reader, limits: ParserLimits) -> ParsedDocument:
    text = _bounded(sanitize_resume_text("\n".join(_pdf_page_text(page) for page in reader.pages)), limits)
    return ParsedDocument(text, "pdf-v4", "good" if text.strip() else "empty")


def _pdf(data: bytes, limits: ParserLimits) -> ParsedDocument:
    from pypdf import PdfReader
    reader = None
    try:
        reader = PdfReader(io.BytesIO(data), strict=True)
        if reader.is_encrypted: raise ParserError("pdf_encrypted")
        if len(reader.pages) > limits.pdf_max_pages: raise ParserError("pdf_page_limit")
    except ParserError: raise
    except Exception:
        reader = None

    try:
        structured = extract_structured_pdf(
            data,
            max_source_bytes=limits.max_source_bytes,
            max_text_chars=limits.max_text_chars,
            max_pages=limits.pdf_max_pages,
        )
        structured = _bounded(sanitize_resume_text(structured), limits)
        if structured.strip():
            return ParsedDocument(structured, "pdf-pdfplumber-v1", "good")
    except StructuredPdfError as error:
        if error.safe_code in {"pdf_encrypted", "pdf_page_limit"}:
            raise ParserError(error.safe_code) from None
    except Exception:
        pass
    if reader is None:
        raise ParserError("pdf_malformed")
    try:
        return _pdf_fallback(reader, limits)
    except ParserError:
        raise
    except Exception:
        raise ParserError("pdf_malformed") from None


def _docx_blocks(document) -> list[str]:
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    values: list[str] = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            value = normalize_resume_line(Paragraph(child, document).text)
            if value:
                values.append(value)
        elif isinstance(child, CT_Tbl):
            table = Table(child, document)
            seen_cells: set[object] = set()
            for row in table.rows:
                for cell in row.cells:
                    cell_key = cell._tc
                    if cell_key in seen_cells:
                        continue
                    seen_cells.add(cell_key)
                    for paragraph in cell.paragraphs:
                        value = normalize_resume_line(paragraph.text)
                        if value:
                            values.append(value)
    return values

def _docx(data: bytes, limits: ParserLimits) -> ParsedDocument:
    try:
        _docx_preflight(data,limits)
        from docx import Document
        document = Document(io.BytesIO(data)); text = _bounded(sanitize_resume_text("\n".join(_docx_blocks(document))), limits)
    except ParserError: raise
    except (zipfile.BadZipFile, KeyError, ValueError): raise ParserError("docx_malformed") from None
    return ParsedDocument(text, "docx-v2", "good" if text.strip() else "empty")

def _txt(data: bytes, limits: ParserLimits) -> ParsedDocument:
    if b"\x00" in data or (data and sum(byte < 9 or 13 < byte < 32 for byte in data) / len(data) > .02): raise ParserError("binary_text_rejected")
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try: text = data.decode(encoding); break
        except UnicodeDecodeError: continue
    else: raise ParserError("text_encoding_unsupported")
    return ParsedDocument(_bounded(text, limits), "txt-v1", "good" if text.strip() else "empty")
