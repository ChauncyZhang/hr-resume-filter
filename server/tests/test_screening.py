import io
import zipfile

import pytest

from server.app.screening.parsers import ParserError, ParserLimits, parse_document
from server.app.screening.rules import ENGINE_VERSION, RuleSnapshot, RuleSnapshotError, score_resume
from server.app.queue.payloads import DEFAULT_PAYLOAD_POLICIES, UnsafePayload
import uuid
from types import SimpleNamespace
from server.app.screening.service import InvalidScreeningTransition, transition_run


def test_rule_engine_matches_legacy_and_returns_structured_ordered_facts() -> None:
    jd = "必须条件：Python, FastAPI, 本科, python\n加分项：Docker, PostgreSQL"
    resume = "5年 Python / FastAPI 后端经验，本科，熟悉 Docker。"
    result = score_resume(resume, RuleSnapshot(jd_text=jd))
    assert ENGINE_VERSION == "rule-v1"
    assert result.score == 92
    assert result.recommendation == "优先沟通"
    assert result.required_hits == ["Python", "FastAPI", "本科"]
    assert result.required_missing == []
    assert result.bonus_hits == ["Docker"]
    assert result.estimated_years == 5


@pytest.mark.parametrize(("jd", "resume", "expected"), [
    ("必须条件：Python, FastAPI, 本科", "3年 Java 专科", (6, "需人工复核", [], ["Python", "FastAPI", "本科"], [], 3)),
    ("required: Python, LLM\nbonus:", "Python 2 years", (38, "需人工复核", ["Python"], ["LLM"], [], 0)),
    ("Python backend distributed systems", "Python backend 4 years", (38, "需人工复核", ["Python", "backend"], ["distributed", "systems"], [], 0)),
    ("硬性要求：中文沟通, 数据分析", "中文沟通，2年数据分析", (79, "可沟通", ["中文沟通", "数据分析"], [], [], 2)),
])
def test_rule_engine_golden_parity_and_missing_cap(jd: str, resume: str, expected: tuple) -> None:
    result = score_resume(resume, RuleSnapshot(jd_text=jd))
    assert (result.score, result.recommendation, result.required_hits, result.required_missing, result.bonus_hits, result.estimated_years) == expected
    if result.required_missing:
        assert result.score <= 59 and result.recommendation == "需人工复核"


@pytest.mark.parametrize(("jd", "resume", "score", "recommendation"), [
    ("required: Python", "Python", 75, "可沟通"),
    ("required: Python\nbonus: Docker", "Python Docker", 90, "优先沟通"),
    ("required: Python", "Python 5年经验", 85, "优先沟通"),
])
def test_rule_weights_are_exactly_75_15_10(jd: str, resume: str, score: int, recommendation: str) -> None:
    result = score_resume(resume, RuleSnapshot(jd_text=jd))
    assert result.score == score
    assert result.recommendation == recommendation


def test_rule_snapshot_overrides_are_immutable_validated_and_change_results() -> None:
    first=RuleSnapshot.from_content("required: Python",{"required_terms":["Python"],"bonus_terms":["Docker"]})
    second=RuleSnapshot.from_content("required: Python",{"required_terms":["Rust"],"bonus_terms":[]})
    assert score_resume("Python Docker 5 years",first).score > score_resume("Python Docker 5 years",second).score
    assert first.required_terms==("Python",) and first.bonus_terms==("Docker",)
    for malformed in (
        {"required_terms":"Python","bonus_terms":[]},
        {"required_terms":["x"*101],"bonus_terms":[]},
        {"required_terms":[],"bonus_terms":[],"unknown":[]},
        {"required_terms":[]},
        {"required_terms":None,"bonus_terms":[]},
        {"must_have":[],"nice_to_have":[],"required_terms":[],"bonus_terms":[]},
    ):
        with pytest.raises(RuleSnapshotError): RuleSnapshot.from_content("required: Python",malformed)


def docx_bytes(text: str) -> bytes:
    from docx import Document
    stream = io.BytesIO(); document = Document(); document.add_paragraph(text); document.save(stream); return stream.getvalue()


def table_docx_bytes() -> bytes:
    from docx import Document
    stream = io.BytesIO(); document = Document(); table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "个人总结"; table.cell(0, 1).text = "负责企业级 AI 平台交付"
    table.cell(1, 0).text = "教育经历"; table.cell(1, 1).text = "2018-2022 某大学 本科"
    document.save(stream); return stream.getvalue()


def pdf_bytes(*, pages: int = 1, encrypted: bool = False) -> bytes:
    from pypdf import PdfWriter
    stream = io.BytesIO(); writer = PdfWriter()
    for _ in range(pages): writer.add_blank_page(width=100, height=100)
    if encrypted: writer.encrypt("secret")
    writer.write(stream); return stream.getvalue()


def test_parser_happy_paths_use_stable_versions_and_quality() -> None:
    txt = parse_document(io.BytesIO("中文 Python 5年".encode("utf-8")), extension=".txt", mime_type="text/plain")
    docx = parse_document(io.BytesIO(docx_bytes("Python 后端")), extension=".docx", mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    pdf = parse_document(io.BytesIO(pdf_bytes()), extension=".pdf", mime_type="application/pdf")
    assert txt.text == "中文 Python 5年" and txt.parser_version == "txt-v1" and txt.quality == "good"
    assert "Python 后端" in docx.text and docx.parser_version == "docx-v2"
    assert pdf.parser_version == "pdf-v4" and pdf.quality == "empty"


def test_docx_parser_reads_table_content_in_document_order() -> None:
    parsed = parse_document(
        io.BytesIO(table_docx_bytes()),
        extension=".docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert parsed.parser_version == "docx-v2"
    assert parsed.text.splitlines() == ["个人总结", "负责企业级 AI 平台交付", "教育经历", "2018-2022 某大学 本科"]


def test_pdf_parser_removes_repeated_standalone_obfuscation_markers(monkeypatch) -> None:
    marker = "bf63fd04e3f2ddac1HJ-3Ni8EFBSwYm9V_6cWOGnn_HZMhll"
    single_identifier = "0123456789abcdef-release_candidate"

    class Page:
        def extract_text(self, visitor_text=None) -> str:
            if visitor_text is not None:
                for index, text in enumerate(("个人简介", "财务系统建设", marker, marker, "项目编号", single_identifier)):
                    visitor_text(text, [1, 0, 0, 1, 0, 0], [1, 0, 0, 1, 40, 40 + index * 20], None, 13)
            return f"个人简介\n财务系统建设\n{marker}\n{marker}\n项目编号\n{single_identifier}"

    class Reader:
        is_encrypted = False
        pages = [Page()]

        def __init__(self, _stream, strict: bool) -> None:
            assert strict is True

    monkeypatch.setattr("pypdf.PdfReader", Reader)

    parsed = parse_document(io.BytesIO(b"%PDF-test"), extension=".pdf", mime_type="application/pdf")

    assert parsed.parser_version == "pdf-v4"
    assert marker not in parsed.text
    assert single_identifier in parsed.text


def test_pdf_parser_reconstructs_two_column_reading_order(monkeypatch) -> None:
    class Page:
        def extract_text(self, visitor_text=None) -> str:
            if visitor_text is not None:
                fragments = (
                    ("个人信息", 40, 220, 18),
                    ("求职意向：财务经理", 40, 250, 13),
                    ("期望城市：深圳", 40, 270, 13),
                    ("个人优势", 258, 57, 18),
                    ("熟练全盘账务处理。", 258, 87, 13),
                    ("工作经历", 258, 206, 18),
                    ("某科技公司 财务经理 2018.08-至今", 258, 237, 13),
                    ("教育经历", 258, 500, 18),
                    ("四川大学 大专 会计 1996-1998", 258, 531, 13),
                )
                for text, x, y, size in fragments:
                    visitor_text(text, [1, 0, 0, 1, 0, 0], [1, 0, 0, 1, x, y], None, size)
            return "个人信息\n求职意向：财务经理\n期望城市：深圳\n熟练全盘账务处理。\n个人优势\n某科技公司 财务经理 2018.08-至今\n工作经历\n四川大学 大专 会计 1996-1998\n教育经历"

    class Reader:
        is_encrypted = False
        pages = [Page()]

        def __init__(self, _stream, strict: bool) -> None:
            assert strict is True

    monkeypatch.setattr("pypdf.PdfReader", Reader)

    parsed = parse_document(io.BytesIO(b"%PDF-test"), extension=".pdf", mime_type="application/pdf")

    assert parsed.text.index("个人优势") < parsed.text.index("熟练全盘账务处理。")
    assert parsed.text.index("工作经历") < parsed.text.index("某科技公司 财务经理 2018.08-至今")
    assert parsed.text.index("教育经历") < parsed.text.index("四川大学 大专 会计 1996-1998")


def test_pdf_parser_uses_plain_order_when_coordinates_are_anomalous(monkeypatch) -> None:
    class MediaBox:
        width = 600
        height = 840

    class Page:
        mediabox = MediaBox()

        def extract_text(self, visitor_text=None) -> str:
            if visitor_text is not None:
                fragments = (
                    ("2024年09月", 0, 0),
                    ("2022年09月", 0, 0),
                    ("2020年09月", 0, 0),
                    ("个人总结", 40, 180),
                    ("具备 Agent 工程经验。", 40, 205),
                    ("教育经历", 40, 900),
                    ("2022年09月-2024年06月 某大学 硕士", 680, 930),
                )
                for text, x, y in fragments:
                    visitor_text(text, [1, 0, 0, 1, 0, 0], [1, 0, 0, 1, x, y], None, 13)
            return "个人总结\n具备 Agent 工程经验。\n教育经历\n2022年09月-2024年06月 某大学 硕士"

    class Reader:
        is_encrypted = False
        pages = [Page()]

        def __init__(self, _stream, strict: bool) -> None:
            assert strict is True

    monkeypatch.setattr("pypdf.PdfReader", Reader)

    parsed = parse_document(io.BytesIO(b"%PDF-test"), extension=".pdf", mime_type="application/pdf")

    assert parsed.text.splitlines()[0] == "个人总结"
    assert "2024年09月2022年09月" not in parsed.text


@pytest.mark.parametrize(("extension", "mime", "data", "code"), [
    (".pdf", "text/plain", b"%PDF-1.4", "file_type_mismatch"),
    (".txt", "text/plain", b"PK\x03\x04junk", "file_magic_mismatch"),
    (".exe", "application/octet-stream", b"MZ", "file_type_not_allowed"),
    (".txt", "text/plain", b"abc\x00def\x00", "binary_text_rejected"),
    (".pdf", "application/pdf", b"not-pdf", "file_magic_mismatch"),
])
def test_parser_rejects_type_magic_and_binary_mismatches(extension: str, mime: str, data: bytes, code: str) -> None:
    with pytest.raises(ParserError) as raised: parse_document(io.BytesIO(data), extension=extension, mime_type=mime)
    assert raised.value.safe_code == code and str(raised.value) == code


def test_pdf_encryption_page_and_malformed_limits_are_typed() -> None:
    with pytest.raises(ParserError) as encrypted: parse_document(io.BytesIO(pdf_bytes(encrypted=True)), extension=".pdf", mime_type="application/pdf")
    assert encrypted.value.safe_code == "pdf_encrypted"
    with pytest.raises(ParserError) as pages: parse_document(io.BytesIO(pdf_bytes(pages=2)), extension=".pdf", mime_type="application/pdf", limits=ParserLimits(pdf_max_pages=1))
    assert pages.value.safe_code == "pdf_page_limit"
    with pytest.raises(ParserError) as malformed: parse_document(io.BytesIO(b"%PDF-broken"), extension=".pdf", mime_type="application/pdf")
    assert malformed.value.safe_code == "pdf_malformed"


def zip_bytes(entries: list[tuple[str, bytes]]) -> bytes:
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in entries: archive.writestr(name, data)
    return stream.getvalue()


@pytest.mark.parametrize(("entries", "limits", "code"), [
    ([('../evil', b'x')], ParserLimits(), "docx_path_traversal"),
    ([('word/vbaProject.bin', b'x')], ParserLimits(), "docx_macro_rejected"),
    ([('a', b'x'), ('b', b'x')], ParserLimits(docx_max_entries=1), "docx_entry_limit"),
    ([('a', b'12345')], ParserLimits(docx_max_uncompressed_bytes=4), "docx_size_limit"),
    ([('a', b'x' * 1000)], ParserLimits(docx_max_compression_ratio=2), "docx_compression_ratio"),
])
def test_docx_zip_preflight_rejections(entries, limits, code: str) -> None:
    with pytest.raises(ParserError) as raised: parse_document(io.BytesIO(zip_bytes(entries)), extension=".docx", mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", limits=limits)
    assert raised.value.safe_code == code


def test_parser_bounds_source_bytes_and_extracted_text_without_leaking_content() -> None:
    secret = "person@example.test resume body"
    with pytest.raises(ParserError) as size: parse_document(io.BytesIO(secret.encode()), extension=".txt", mime_type="text/plain", limits=ParserLimits(max_source_bytes=4))
    assert size.value.safe_code == "file_too_large" and secret not in str(size.value)
    with pytest.raises(ParserError) as text: parse_document(io.BytesIO(secret.encode()), extension=".txt", mime_type="text/plain", limits=ParserLimits(max_text_chars=4))
    assert text.value.safe_code == "text_limit_exceeded" and secret not in str(text.value)


def test_future_screening_queue_payloads_are_registered_and_opaque_only() -> None:
    item_id=str(uuid.uuid4()); org_id=str(uuid.uuid4()); jd_id=str(uuid.uuid4()); rule_id=str(uuid.uuid4()); result_id=str(uuid.uuid4()); config_id=str(uuid.uuid4()); prompt_id=str(uuid.uuid4())
    assert DEFAULT_PAYLOAD_POLICIES.validate_job("screening.parse_item", {"organization_id":org_id,"screening_item_id":item_id,"parser_version":"parser-v1"})["screening_item_id"] == item_id
    assert DEFAULT_PAYLOAD_POLICIES.validate_job("screening.score_item", {"organization_id":org_id,"screening_item_id":item_id,"jd_version_id":jd_id,"rule_version_id":rule_id,"rule_engine_version":"rule-v1"})["rule_engine_version"] == "rule-v1"
    llm_payload={"organization_id":org_id,"screening_item_id":item_id,"screening_result_id":result_id,"config_id":config_id,"config_version":1,"prompt_version_id":prompt_id}
    assert DEFAULT_PAYLOAD_POLICIES.validate_job("screening.llm_score_item",llm_payload)==llm_payload
    with pytest.raises(UnsafePayload): DEFAULT_PAYLOAD_POLICIES.validate_job("screening.parse_item", {"organization_id":org_id,"screening_item_id":item_id,"parser_version":"parser-v1","resume_text":"secret"})
    with pytest.raises(UnsafePayload): DEFAULT_PAYLOAD_POLICIES.validate_job("screening.llm_score_item",{**llm_payload,"resume_text":"secret"})


def test_run_transition_contract_supports_rule_only_and_future_llm_paths() -> None:
    direct = SimpleNamespace(status="rule_scoring", version=1)
    transition_run(direct, "completed"); assert direct.status == "completed" and direct.version == 2
    llm = SimpleNamespace(status="rule_scoring", version=3)
    transition_run(llm, "llm_scoring"); transition_run(llm, "partial")
    assert llm.status == "partial" and llm.version == 5
    for terminal in ("completed", "partial", "failed", "cancelled"):
        run = SimpleNamespace(status="llm_scoring", version=1); transition_run(run, terminal); assert run.status == terminal
    with pytest.raises(InvalidScreeningTransition): transition_run(SimpleNamespace(status="queued", version=1), "llm_scoring")
